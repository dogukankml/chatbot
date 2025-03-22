import streamlit as st
from openai import OpenAI
from dotenv import load_dotenv
import os
import httpx
import time
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import mimetypes
import PyPDF2
import docx2txt

# Load environment variables
load_dotenv()

# Configure OpenAI with custom HTTP client
http_client = httpx.Client()
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
    http_client=http_client
)

# Configure Streamlit page
st.set_page_config(
    page_title="Chatbot",
    page_icon="🤖",
    layout="centered"
)

# Add custom CSS
st.markdown("""
    <style>
    .stTextInput > div > div > input {
        background-color: #f0f2f6;
    }
    .stButton > button {
        border-radius: 4px;
        padding: 2px 10px;
        font-size: 14px;
        background-color: #ff4b4b;
        color: white;
        border: none;
        width: auto !important;
    }
    .stButton > button:hover {
        background-color: #ff3333;
    }
    .stTab {
        background-color: #f0f2f6;
        border-radius: 10px;
        padding: 10px;
        margin: 5px;
    }
    /* Reverse the chat messages container */
    .main > div:has(.stChatMessage) {
        flex-direction: column;
    }
    /* Style for clear button container */
    [data-testid="stHorizontalBlock"] {
        position: fixed !important;
        bottom: 40px !important;
        right: 20px !important;
        width: 478px !important;
        z-index: 1000;
        background: transparent;
    }
    [data-testid="stHorizontalBlock"] > div:first-child {
        display: none !important;
    }
    [data-testid="stHorizontalBlock"] > div:last-child {
        width: auto !important;
        flex: 0 1 auto !important;
        padding: 0 !important;
    }
    [data-testid="stHorizontalBlock"] button {
        margin: 0 !important;
    }
    /* Typing animation */
    @keyframes typing {
        0% { content: ""; }
        25% { content: "."; }
        50% { content: ".."; }
        75% { content: "..."; }
        100% { content: ""; }
    }
    .typing::after {
        content: "";
        animation: typing 2s infinite;
        display: inline-block;
        width: 20px;
        text-align: left;
    }
    /* New Chat button styles */
    [data-testid="baseButton-primary"] {
        position: fixed !important;
        top: 70px !important;
        right: 20px !important;
        z-index: 1000 !important;
        background-color: #4CAF50 !important;
        color: white !important;
        border: none !important;
        padding: 5px 15px !important;
        border-radius: 4px !important;
    }
    [data-testid="baseButton-primary"]:hover {
        background-color: #45a049 !important;
        color: white !important;
        border: none !important;
    }
    /* Tab container styles */
    .stTabs [data-baseweb="tab-list"] {
        gap: 2px;
    }
    .stTabs [data-baseweb="tab"] {
        position: relative;
        padding-right: 20px !important;
    }
    </style>
    """, unsafe_allow_html=True)

# Function to export chat as TXT
def export_chat_txt(messages, tab_name):
    chat_text = [f"Chat Export - {tab_name}"]
    chat_text.append(f"Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    
    for msg in messages:
        role = msg["role"].capitalize()
        content = msg["content"]
        chat_text.append(f"{role}: {content}\n")
    
    return "\n".join(chat_text)

# Function to export chat as PDF
def export_chat_pdf(messages, tab_name):
    # Create a BytesIO buffer to receive PDF data
    buffer = BytesIO()
    
    # Create the PDF document
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Create custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30
    )
    timestamp_style = ParagraphStyle(
        'Timestamp',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.gray,
        spaceAfter=20
    )
    message_style = ParagraphStyle(
        'Message',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=15,
        leading=14
    )
    
    # Add title
    story.append(Paragraph(f"Chat Export - {tab_name}", title_style))
    story.append(Paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", timestamp_style))
    story.append(Spacer(1, 20))
    
    # Add messages
    for msg in messages:
        role = msg["role"].capitalize()
        content = msg["content"].replace('\n', '<br/>')  # Convert newlines to HTML breaks
        story.append(Paragraph(
            f"<b>{role}:</b> {content}",
            message_style
        ))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

# Function to export chat as DOCX
def export_chat_docx(messages, tab_name):
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Chat Export - {tab_name}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph()
    timestamp_run = timestamp.add_run(f'Exported on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    timestamp_run.font.size = Pt(10)
    timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add some space
    doc.add_paragraph()
    
    # Add messages
    for msg in messages:
        # Add message
        paragraph = doc.add_paragraph()
        role = msg["role"].capitalize()
        role_run = paragraph.add_run(f'{role}: ')
        role_run.bold = True
        content_run = paragraph.add_run(msg["content"])
        
        # Add spacing after message
        paragraph.spacing_after = Pt(12)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to read file content based on file type
def read_file_content(uploaded_file):
    file_type = mimetypes.guess_type(uploaded_file.name)[0]
    content = ""
    
    try:
        if file_type == "text/plain":
            content = uploaded_file.getvalue().decode()
        elif file_type == "application/pdf":
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            content = docx2txt.process(uploaded_file)
        else:
            content = f"File type {file_type} is not supported. Supported types are: txt, pdf, doc, docx"
    except Exception as e:
        content = f"Error reading file: {str(e)}"
    
    return content

# Initialize session state for tabs
if "tabs" not in st.session_state:
    st.session_state.tabs = {
        "Chat 1": [{"role": "assistant", "content": "How can I help you?"}]
    }
if "tab_counter" not in st.session_state:
    st.session_state.tab_counter = 1
if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = {}

# Function to create new chat tab
def create_new_chat():
    st.session_state.tab_counter += 1
    new_tab_name = f"Chat {st.session_state.tab_counter}"
    
    while new_tab_name in st.session_state.tabs:
        st.session_state.tab_counter += 1
        new_tab_name = f"Chat {st.session_state.tab_counter}"

    st.session_state.tabs[new_tab_name] = [{"role": "assistant", "content": "How can I help you?"}]
    st.session_state.current_tab = new_tab_name


# Display chat header and new chat button
st.title("🤖 Chatbot")

# Add New Chat button
if st.button("New Chat", type="primary", key="new_chat_button"):
    create_new_chat()
    st.rerun()

# Create tabs dynamically
tabs = st.tabs(list(st.session_state.tabs.keys()))

# Function to get all messages from other tabs
def get_other_tabs_context(current_tab):
    other_messages = []
    for tab, messages in st.session_state.tabs.items():
        if tab != current_tab and messages:
            other_messages.append(f"Messages from {tab}:")
            for msg in messages:
                other_messages.append(f"{msg['role']}: {msg['content']}")
    return "\n".join(other_messages)

# Function to clear chat history
def clear_chat(tab_name):
    st.session_state.tabs[tab_name] = [{"role": "assistant", "content": "How can I help you?"}]
    st.rerun()

def close_chat(tab_name):
    if tab_name in st.session_state.tabs:
        del st.session_state.tabs[tab_name]

        remaining_tabs = list(st.session_state.tabs.keys())
        st.session_state.current_tab = remaining_tabs[-1] if remaining_tabs else "Chat 1"

        updated_tabs = {}
        for i, old_tab_name in enumerate(remaining_tabs, start=1):
            updated_tab_name = f"Chat {i}"
            updated_tabs[updated_tab_name] = st.session_state.tabs[old_tab_name]

        st.session_state.tabs = updated_tabs
        st.session_state.tab_counter = len(updated_tabs)  
        
        st.rerun()


# Function to handle chat interaction
def handle_chat(tab_name, tab):
    with tab:
        # Only show file uploader in the first tab
        if tab_name == "Chat 1":
            st.sidebar.markdown("### Upload Files")
            uploaded_file = st.sidebar.file_uploader(
                "Upload a file to discuss",
                type=["txt", "pdf", "doc", "docx"],
                key=f"file_uploader_{tab_name}",
                accept_multiple_files=True  # Enable multiple file upload
            )
            
            # Handle file upload silently
            if uploaded_file is not None:
                for file in uploaded_file:
                    if file.name not in st.session_state.uploaded_files:
                        file_content = read_file_content(file)
                        st.session_state.uploaded_files[file.name] = file_content
                        st.rerun()
        
        # Create a container for chat messages
        chat_container = st.container()
        
        with chat_container:
            # Display chat history for current tab
            for message in st.session_state.tabs[tab_name]:
                if message["role"] != "system":  # Skip system messages
                    with st.chat_message(message["role"]):
                        st.markdown(message["content"])
        
        # Add buttons container at the bottom right
        col1, col2, col3, col4, col5 = st.columns([3, 1, 1.2, 1, 1])

        with col2:
            if st.button("Clear Chat", key=f"clear_{tab_name}", help="Clear chat history"):
                clear_chat(tab_name)

        # İndirme butonlarının state'lerini başlat
        if "download_states" not in st.session_state:
            st.session_state.download_states = {}

        # Butonları aç/kapa fonksiyonu
        def toggle_download_button(tab_name, file_type):
            key = f"download_{file_type}_{tab_name}"
            st.session_state.download_states[key] = not st.session_state.download_states.get(key, False)

        # Export buttons
        with col3:
            if st.button("Export DOCX", key=f"docx_{tab_name}", help="Export chat as DOCX", type="secondary"):
                toggle_download_button(tab_name, "docx")

            if st.session_state.download_states.get(f"download_docx_{tab_name}", False):
                docx_buffer = export_chat_docx(st.session_state.tabs[tab_name], tab_name)
                st.download_button(
                    label="Download DOCX",
                    data=docx_buffer,
                    file_name=f"chat_export_{tab_name.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_docx_{tab_name}"
                )

        with col4:
            if st.button("Export TXT", key=f"txt_{tab_name}", help="Export chat as TXT", type="secondary"):
                toggle_download_button(tab_name, "txt")

            if st.session_state.download_states.get(f"download_txt_{tab_name}", False):
                txt_data = export_chat_txt(st.session_state.tabs[tab_name], tab_name)
                st.download_button(
                    label="Download TXT",
                    data=txt_data,
                    file_name=f"chat_export_{tab_name.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    key=f"download_txt_{tab_name}"
                )

        with col5:
            if st.button("Export PDF", key=f"pdf_{tab_name}", help="Export chat as PDF", type="secondary"):
                toggle_download_button(tab_name, "pdf")

            if st.session_state.download_states.get(f"download_pdf_{tab_name}", False):
                pdf_buffer = export_chat_pdf(st.session_state.tabs[tab_name], tab_name)
                st.download_button(
                    label="Download PDF",
                    data=pdf_buffer,
                    file_name=f"chat_export_{tab_name.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf",
                    key=f"download_pdf_{tab_name}"
                )
 
        # Create a container for the input at the bottom
        with st.container():
            st.markdown('<div class="chat-input">', unsafe_allow_html=True)
            # Get user input with unique key for each tab
            if prompt := st.chat_input("Type your message here...", key=f"chat_input_{tab_name}"):
                # Add user message to chat history
                st.session_state.tabs[tab_name].append({"role": "user", "content": prompt})
                
                try:
                    # Get context from other tabs
                    other_tabs_context = get_other_tabs_context(tab_name)
                    
                    # Prepare messages with context and file information
                    system_message = "You are a helpful assistant. "
                    
                    # Add file contents to system message
                    if st.session_state.uploaded_files:
                        system_message += "\n\nAvailable files and their contents:\n"
                        for filename, content in st.session_state.uploaded_files.items():
                            system_message += f"\nFile: {filename}\nContent:\n{content}\n---\n"
                    
                    if other_tabs_context:
                        system_message += f"\n\nContext from other conversations:\n{other_tabs_context}"
                    
                    messages = [
                        {"role": "system", "content": system_message}
                    ]
                    messages.extend([
                        {"role": m["role"], "content": m["content"]}
                        for m in st.session_state.tabs[tab_name]
                    ])
                    
                    # Show typing animation
                    with st.chat_message("assistant"):
                        st.markdown('<div class="typing">Thinking</div>', unsafe_allow_html=True)
                        
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",  
                            messages=messages,
                            stream=True
                        )
                        
                        full_response = ""
                        message_placeholder = st.empty()
                        
                        # Stream the response
                        for chunk in response:
                            if chunk.choices[0].delta.content is not None:
                                full_response += chunk.choices[0].delta.content
                                message_placeholder.markdown(full_response)
                                time.sleep(0.01)  # Small delay for smoother animation
                    
                    # Add assistant response to chat history
                    st.session_state.tabs[tab_name].append({"role": "assistant", "content": full_response})
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")
                    if "Invalid API key" in str(e):
                        st.warning("Please enter a valid OpenAI API key in the .env file")
            st.markdown('</div>', unsafe_allow_html=True)


# Handle each tab
for tab_name, tab in zip(list(st.session_state.tabs.keys()), tabs):
    with tab:
        # Sekme başlığı ve kapatma butonu
        st.markdown(f"### {tab_name}")
        if tab_name != "Chat 1":
            if st.button(f"Close {tab_name}  ✘ ", key=f"close_{tab_name}"):
                close_chat(tab_name)
        handle_chat(tab_name, tab)


